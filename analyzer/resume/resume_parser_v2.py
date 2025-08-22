"""
改进的简历解析器，支持多种文档格式
"""
import os
import re
import logging
from io import BytesIO

# 文档解析库
try:
    import docx
except ImportError:
    docx = None
    
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    
try:
    from pdfplumber import PDF
except ImportError:
    PDF = None


class ResumeParserV2:
    """改进的简历文件解析器"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
    def parse_uploaded_file(self, file_obj):
        """解析上传的文件对象，返回纯文本内容"""
        filename = file_obj.filename
        file_ext = os.path.splitext(filename)[1].lower()
        
        self.logger.info(f"开始解析文件: {filename}, 类型: {file_ext}")
        
        # 读取文件内容
        file_content = file_obj.read()
        file_obj.seek(0)  # 重置文件指针
        
        try:
            if file_ext == '.pdf':
                return self._parse_pdf_bytes(file_content)
            elif file_ext in ['.docx', '.doc']:
                return self._parse_docx_bytes(file_content)
            elif file_ext == '.txt':
                return self._parse_txt_bytes(file_content)
            else:
                # 尝试作为文本文件处理
                return self._parse_txt_bytes(file_content)
                
        except Exception as e:
            self.logger.error(f"文件解析失败: {e}")
            # 返回错误信息而不是抛出异常
            return f"文件解析失败: {str(e)}"
    
    def _parse_pdf_bytes(self, file_content):
        """解析PDF字节流"""
        text_content = ""
        
        # 优先使用pdfplumber（更准确）
        if PDF:
            try:
                with PDF(BytesIO(file_content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
                return text_content.strip()
            except Exception as e:
                self.logger.warning(f"pdfplumber解析失败: {e}")
        
        # 回退到PyPDF2
        if PyPDF2:
            try:
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_content += text + "\n"
                return text_content.strip()
            except Exception as e:
                self.logger.error(f"PyPDF2解析失败: {e}")
                
        return "PDF解析失败：缺少必要的库"
    
    def _parse_docx_bytes(self, file_content):
        """解析DOCX字节流"""
        if not docx:
            return "DOCX解析失败：缺少python-docx库"
            
        try:
            # 创建内存中的文件对象
            doc_stream = BytesIO(file_content)
            doc = docx.Document(doc_stream)
            
            text_content = ""
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
                        
            return text_content.strip()
            
        except Exception as e:
            self.logger.error(f"DOCX解析错误: {e}")
            # 如果是doc格式（老版本），尝试其他方法
            if '.doc' in str(e).lower():
                return "暂不支持DOC格式，请转换为DOCX或PDF格式"
            return f"DOCX解析失败: {str(e)}"
    
    def _parse_txt_bytes(self, file_content):
        """解析文本字节流"""
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'big5']
        
        for encoding in encodings:
            try:
                return file_content.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
                
        # 如果都失败，使用错误处理
        return file_content.decode('utf-8', errors='ignore').strip()