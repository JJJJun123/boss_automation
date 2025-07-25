import os
import re
import docx
import PyPDF2
from io import BytesIO


class ResumeParser:
    """简历文件解析器，支持PDF、DOCX、TXT格式"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def parse_file(self, file_path):
        """解析简历文件，返回文本内容"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext == '.docx':
            return self._parse_docx(file_path)
        elif file_ext == '.txt':
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    def parse_uploaded_file(self, file_obj):
        """解析上传的文件对象"""
        filename = file_obj.filename
        file_ext = os.path.splitext(filename)[1].lower()
        
        # 读取文件内容到内存
        file_content = file_obj.read()
        file_obj.seek(0)  # 重置文件指针
        
        if file_ext == '.pdf':
            return self._parse_pdf_from_bytes(file_content)
        elif file_ext == '.docx':
            return self._parse_docx_from_bytes(file_content)
        elif file_ext == '.txt':
            return self._parse_txt_from_bytes(file_content)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    def _parse_pdf(self, file_path):
        """解析PDF文件"""
        try:
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            return text_content.strip()
        except Exception as e:
            raise Exception(f"PDF解析失败: {str(e)}")
    
    def _parse_pdf_from_bytes(self, file_content):
        """从字节流解析PDF"""
        try:
            text_content = ""
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            return text_content.strip()
        except Exception as e:
            raise Exception(f"PDF解析失败: {str(e)}")
    
    def _parse_docx(self, file_path):
        """解析DOCX文件"""
        try:
            doc = docx.Document(file_path)
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            return text_content.strip()
        except Exception as e:
            raise Exception(f"DOCX解析失败: {str(e)}")
    
    def _parse_docx_from_bytes(self, file_content):
        """从字节流解析DOCX"""
        try:
            doc = docx.Document(BytesIO(file_content))
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            return text_content.strip()
        except Exception as e:
            raise Exception(f"DOCX解析失败: {str(e)}")
    
    def _parse_txt(self, file_path):
        """解析TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(file_path, 'r', encoding='gbk') as file:
                return file.read().strip()
        except Exception as e:
            raise Exception(f"TXT解析失败: {str(e)}")
    
    def _parse_txt_from_bytes(self, file_content):
        """从字节流解析TXT"""
        try:
            return file_content.decode('utf-8').strip()
        except UnicodeDecodeError:
            try:
                return file_content.decode('gbk').strip()
            except Exception as e:
                raise Exception(f"TXT解析失败: {str(e)}")
    
    def extract_basic_info(self, text_content):
        """从文本中提取基本信息"""
        info = {
            'name': self._extract_name(text_content),
            'phone': self._extract_phone(text_content),
            'email': self._extract_email(text_content),
            'experience_years': self._extract_experience_years(text_content),
            'education': self._extract_education(text_content),
            'skills': self._extract_skills(text_content)
        }
        return info
    
    def _extract_name(self, text):
        """提取姓名"""
        # 简单的姓名提取逻辑
        lines = text.split('\n')
        for line in lines[:5]:  # 通常姓名在前几行
            line = line.strip()
            if len(line) >= 2 and len(line) <= 10 and re.match(r'^[\u4e00-\u9fa5]{2,4}$', line):
                return line
        return "未知"
    
    def _extract_phone(self, text):
        """提取电话号码"""
        phone_pattern = r'1[3-9]\d{9}'
        match = re.search(phone_pattern, text)
        return match.group(0) if match else "未提供"
    
    def _extract_email(self, text):
        """提取邮箱"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(email_pattern, text)
        return match.group(0) if match else "未提供"
    
    def _extract_experience_years(self, text):
        """提取工作年限"""
        # 查找年限相关的表述
        patterns = [
            r'(\d+)年.*?经验',
            r'工作.*?(\d+)年',
            r'从业.*?(\d+)年',
            r'(\d+)年.*?工作'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return int(match.group(1))
        return 0
    
    def _extract_education(self, text):
        """提取教育背景"""
        education_keywords = ['本科', '硕士', '博士', '学士', '研究生', '大学']
        
        for keyword in education_keywords:
            if keyword in text:
                return keyword
        return "未知"
    
    def _extract_skills(self, text):
        """提取技能关键词"""
        # 常见技能关键词
        skill_keywords = [
            'Python', 'Java', 'JavaScript', 'SQL', 'R', 'C++', 'C#',
            '数据分析', '机器学习', '深度学习', '人工智能', 'AI',
            '风险管理', '风险控制', '量化分析', '建模',
            'Excel', 'PPT', 'Word', 'SPSS', 'SAS', 'Tableau',
            '项目管理', '团队管理', '沟通能力', '领导力'
        ]
        
        found_skills = []
        text_upper = text.upper()
        
        for skill in skill_keywords:
            if skill.upper() in text_upper or skill in text:
                found_skills.append(skill)
        
        return found_skills[:10]  # 返回前10个技能