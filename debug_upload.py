#!/usr/bin/env python3
"""
调试简历上传问题
"""
import requests
import json
from io import BytesIO

# 测试1: 检查API是否可达
try:
    print("测试1: 检查API健康状态...")
    response = requests.get('http://localhost:5000/api/health')
    print(f"状态码: {response.status_code}")
    print(f"响应: {response.json()}")
except Exception as e:
    print(f"API健康检查失败: {e}")
    exit(1)

# 测试2: 测试简历上传
try:
    print("\n测试2: 上传简单文本文件...")
    
    # 创建一个简单的文本文件
    text_content = """张三
电话: 13812345678
邮箱: zhangsan@example.com

工作经验:
2020-2023 ABC公司 - 数据分析师
负责数据分析和建模工作

教育背景:
2016-2020 某大学 - 计算机科学

技能:
Python, SQL, 机器学习"""
    
    files = {'resume': ('test_resume.txt', text_content, 'text/plain')}
    
    print("正在发送请求...")
    response = requests.post('http://localhost:5000/api/upload_resume', files=files, timeout=120)
    
    print(f"响应状态码: {response.status_code}")
    print(f"响应头: {dict(response.headers)}")
    
    if response.status_code == 200:
        try:
            data = response.json()
            print("✅ 解析JSON成功")
            print(f"成功状态: {data.get('success')}")
            if data.get('success'):
                print(f"姓名: {data.get('resume_data', {}).get('name')}")
                print(f"AI评分: {data.get('ai_analysis', {}).get('competitiveness_score')}")
            else:
                print(f"上传失败: {data.get('error')}")
        except json.JSONDecodeError as e:
            print(f"❌ JSON解析失败: {e}")
            print(f"原始响应: {response.text[:500]}...")
    else:
        print(f"❌ HTTP错误: {response.status_code}")
        print(f"错误响应: {response.text}")
        
except requests.exceptions.Timeout:
    print("❌ 请求超时")
except requests.exceptions.ConnectionError:
    print("❌ 连接错误")
except Exception as e:
    print(f"❌ 其他错误: {e}")
    import traceback
    traceback.print_exc()

# 测试3: 测试DOCX文件
try:
    print("\n测试3: 上传DOCX文件...")
    import docx
    
    # 创建简单的DOCX文档
    doc = docx.Document()
    doc.add_heading('李四', 0)
    doc.add_paragraph('电话: 13987654321')
    doc.add_paragraph('邮箱: lisi@example.com')
    doc.add_paragraph('工作经验: 5年软件开发经验')
    
    # 保存到内存
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    files = {'resume': ('test_resume.docx', docx_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
    
    response = requests.post('http://localhost:5000/api/upload_resume', files=files, timeout=120)
    
    print(f"DOCX响应状态码: {response.status_code}")
    
    if response.status_code == 200:
        data = response.json()
        if data.get('success'):
            print("✅ DOCX上传成功")
            print(f"姓名: {data.get('resume_data', {}).get('name')}")
        else:
            print(f"❌ DOCX上传失败: {data.get('error')}")
    
except Exception as e:
    print(f"DOCX测试失败: {e}")